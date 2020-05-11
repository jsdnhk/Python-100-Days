"""
讀取Word文件

Version: 0.1
Author: 駱昊
Date: 2018-03-26
"""

from docx import Document

doc = Document('./res/用函數還是用複雜的表達式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].runs[0].text)

content = []
for para in doc.paragraphs:
    content.append(para.text)
print(''.join(content))
